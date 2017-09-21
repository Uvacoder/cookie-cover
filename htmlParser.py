#Parses the Recruit Guelph page given the URL
from selenium import webdriver
from bs4 import BeautifulSoup
import re
import json
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DEBUG = False

#Log printing
def log(d):
    if DEBUG:
        print(d)

#Parsing for the address portion of the cover letter
def parse(driver):
    #Get the job page
    html = driver.page_source
    #Remove all newlines, tabs and funky music
    html = html.replace("\t", "")
    html = html.replace("\n", "")
    html = html.replace("–", "-")
    html = html.replace("’", "'")
    html = html.replace("‘", "'")
    html = html.replace("—", "-")
    html = html.replace("●", "")
    html = html.replace("“", "\"")
    html = html.replace("”", "\"")
    html = html.replace("​", " ")

    log(html)
    bs = BeautifulSoup(html, "html.parser")
    #Find all tables
    tableList = bs.find_all("table")
    tdList = tableList[3].findChildren('td', {"width" : "75%"})
    for td in tdList:
        log(td.get_text())
    #Find job title
    h1List = bs.find_all("h1")
    for h1 in h1List:
        log(h1.get_text())
    #Adding job title to tdList
    tdList.extend(h1List[1])
    return tdList

def documentWrite(tdList):
    #Fix list length if website is not included in application
    if(len(tdList) != 14 ):
        log("Fixing list")
        tdList.insert(7, "Website")
    #Open document
    coverLetter = Document('test.docx')
    #Create cover letter header
    headerTable = coverLetter.add_table(rows = 0, cols = 1, style = "CoverLetter")
    row1 = headerTable.add_row().cells
    headerName = row1[0].text = "Bradley Leonard (4 Month Co-op)"
    row1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #row1[0].paragraphs[0].font.size = Pt(24)
    row2 = headerTable.add_row().cells
    headerBody = row2[0].text = "700 Bartlett Drive, Labrador City, NL A2V 1G8 (permanent)\n1 Hales Cres. Unit 503, Guelph, ON N1G 1P4 (current)\nCell: (709) 280-6600 Alt: (709) 944-7797 Email: bleona02@uoguelph.ca"
    row2[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #Write date letter was created
    log("Current time: " + time.strftime("%B %d, %Y"))
    date = coverLetter.add_paragraph("\n" + time.strftime("%B %d, %Y"))
    #Write address and other employer information
    address = coverLetter.add_paragraph(tdList[4].get_text() + " " + tdList[5].get_text() + "\n")
    address.add_run(tdList[6].get_text() + "\n")
    address.add_run(tdList[1].get_text() + "\n")
    address.add_run(tdList[8].get_text() + "\n")
    address.add_run(tdList[9].get_text() + ", " + tdList[10].get_text() + " " + tdList[11].get_text() + "\n")
    #Write greeting
    greeting = coverLetter.add_paragraph("Dear " + tdList[3].get_text() + " " + tdList[5].get_text() + ",")
    #Write self introduction and interest
    body1 = coverLetter.add_paragraph("My name is Bradley Leonard and I am a 3rd year Engineering Systems and Computing student at the University of Guelph. With my passion for both programming and robotics I am aiming for the Mechatronics stream of my program. ")
    body1.paragraph_format.first_line_indent = Inches(0.25)
    body1.add_run("I am currently looking for a 4 month co-op position in the winter of 2018. Your posted position on Recruit Guelph for " + tdList[13] + " has acquired my attention as I possess a fascination with coding in a multitude of languages and a love for seeing my creations come to life right before my eyes.")
    #Write tech introduction
    body2 = coverLetter.add_paragraph("I am a hardworking individual with a passion for technology and am quick to pick up additional skills either through self-teaching or being taught by others. This can be shown through my co-development of the app “Plannit” or my co-founding of FreebieMapp. They are two broad projects that exhibit my craving for a challenge and yearning to learn. I look to do the best job I can in the most efficient manner possible. In addition, I have acquired a large set of skills presently and am always willing to expand this list.")
    body2.paragraph_format.first_line_indent = Inches(0.25)
    #Write personal skills
    body3 = coverLetter.add_paragraph("I am also a friendly, patient and determined individual. These interpersonal skills come from my vast experience with playing and coaching sports along with many other types of social events. However when the task requires it, I can also work autonomously.")
    body3.paragraph_format.first_line_indent = Inches(0.25)
    #Write conclusion
    body4 = coverLetter.add_paragraph("I would welcome the opportunity to discuss your needs and my qualifications in more details.  Thank you for your consideration and time, I can be reached at any of the mentioned contacts.")
    body4.paragraph_format.first_line_indent = Inches(0.25)
    #Write closing
    closing = coverLetter.add_paragraph("Respectfully yours,\nBradley Leonard")
    #Save document
    coverLetter.save("Bradley Leonard Cover Letter " + tdList[0].get_text() + " " + tdList[13] + ".docx")

#Load up configs
with open('config.json') as configs:
    jsonConfigs = json.load(configs)
    log(jsonConfigs)
#Start up Chrome and get user to navigate to page
print("Starting Chrome Window...")
print("Please navigate to the job page.")
driver = webdriver.Chrome("C:/Users/Bradley/AppData/Local/Programs/Python/Python35/Scripts/chromedriver.exe")
driver.get("https://www.recruitguelph.ca/students/student-login.htm")
#Log into SSO
username = driver.find_element_by_id("inputUsername")
password = driver.find_element_by_id("inputPassword")
username.send_keys(jsonConfigs["username"])
password.send_keys(jsonConfigs["password"])
driver.find_element_by_class_name("btn").click()
#Navigate to job pages
driver.find_element_by_link_text("Co-op").click()
driver.find_element_by_link_text("Co-op Job Postings").click()

#Wait for user to be ready
ready = False
while(ready == False):
    readyAnswer = input("Are you ready? ")
    if(readyAnswer == "y" or readyAnswer == "yes"):
        ready = True
tdList = parse(driver)
documentWrite(tdList)
